"""Document loaders for various file formats."""

import os
from pathlib import Path
from typing import Optional
from datetime import datetime
import logging

from .models import Document, DocumentMetadata, DocumentType

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load documents from various file formats."""
    
    def __init__(self):
        """Initialize document loader."""
        self._loaders = {
            ".txt": self._load_text,
            ".md": self._load_markdown,
            ".pdf": self._load_pdf,
            ".html": self._load_html,
            ".htm": self._load_html,
            ".docx": self._load_docx,
        }
    
    def load(self, file_path: str, **metadata_overrides) -> Document:
        """
        Load a document from a file.
        
        Args:
            file_path: Path to the document file
            **metadata_overrides: Additional metadata to override/add
            
        Returns:
            Document object with content and metadata
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type
        extension = path.suffix.lower()
        if extension not in self._loaders:
            logger.warning(f"Unsupported file type: {extension}, treating as text")
            extension = ".txt"
        
        # Load content
        content = self._loaders[extension](path)
        
        # Create metadata
        doc_type = self._get_document_type(extension)
        stat = path.stat()
        
        metadata = DocumentMetadata(
            source=str(path.absolute()),
            document_type=doc_type,
            title=metadata_overrides.get("title", path.stem),
            file_size=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            tags=metadata_overrides.get("tags", []),
            category=metadata_overrides.get("category"),
            custom_metadata=metadata_overrides.get("custom_metadata", {}),
        )
        
        return Document(content=content, metadata=metadata)
    
    def _get_document_type(self, extension: str) -> DocumentType:
        """Map file extension to document type."""
        mapping = {
            ".txt": DocumentType.TEXT,
            ".md": DocumentType.MARKDOWN,
            ".pdf": DocumentType.PDF,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".docx": DocumentType.DOCX,
        }
        return mapping.get(extension, DocumentType.UNKNOWN)
    
    def _load_text(self, path: Path) -> str:
        """Load plain text file."""
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 if utf-8 fails
            with open(path, "r", encoding="latin-1") as f:
                return f.read()
    
    def _load_markdown(self, path: Path) -> str:
        """Load markdown file."""
        return self._load_text(path)
    
    def _load_pdf(self, path: Path) -> str:
        """Load PDF file using PyPDF2."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF support. "
                "Install it with: pip install PyPDF2"
            )
        
        text_content = []
        with open(path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        return "\n\n".join(text_content)
    
    def _load_html(self, path: Path) -> str:
        """Load HTML file and extract text."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML support. "
                "Install it with: pip install beautifulsoup4"
            )
        
        html_content = self._load_text(path)
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _load_docx(self, path: Path) -> str:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX support. "
                "Install it with: pip install python-docx"
            )
        
        doc = DocxDocument(path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n\n".join(paragraphs)
